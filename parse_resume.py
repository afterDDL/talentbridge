import json
import pathlib
import sys


def parse_pdf(file_path):
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def parse_docx(file_path):
    from docx import Document

    document = Document(str(file_path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            paragraphs.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(paragraphs)


def parse_text(file_path):
    return file_path.read_text(encoding="utf-8", errors="replace")


def main():
    if len(sys.argv) != 2:
        raise ValueError("需要提供文件路径")

    file_path = pathlib.Path(sys.argv[1]).resolve()
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        text = parse_pdf(file_path)
    elif suffix == ".docx":
        text = parse_docx(file_path)
    elif suffix in {".txt", ".md"}:
        text = parse_text(file_path)
    else:
        raise ValueError(f"不支持的文件格式：{suffix}")

    cleaned = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if len(cleaned) < 20:
        raise ValueError("未提取到足够文本，文件可能是扫描件或内容为空")

    print(json.dumps({"text": cleaned}, ensure_ascii=True))


if __name__ == "__main__":
    try:
        main()
    except Exception as error:
        print(json.dumps({"error": str(error)}, ensure_ascii=True))
        sys.exit(1)
